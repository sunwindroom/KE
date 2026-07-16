"""从已上传文档的原始字节中抽取纯文本，供知识抽取环节使用。

支持：PDF（含多页拼接）、DOCX（段落 + 表格）、纯文本/CSV（按 UTF-8/GBK 尝试解码）。
不支持：扫描件 OCR（需要额外的 OCR 引擎，超出当前范围）——遇到无法提取文本的
扫描版 PDF 时会返回空字符串并记录警告，调用方应据此提示用户改用可编辑文档或
人工录入，而不是假装抽到了内容。
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx_text(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def extract_text(content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            text = _extract_pdf_text(content)
            if not text.strip():
                logger.warning("PDF %s 未能提取到文本，可能是扫描件（暂不支持 OCR）", filename)
            return text
        if ext == "docx":
            return _extract_docx_text(content)
        if ext in ("txt", "csv", "md"):
            return _extract_plain_text(content)
        logger.warning("不支持自动解析的文件类型: %s，跳过文本抽取", ext)
        return ""
    except Exception:
        logger.exception("解析文档 %s 失败", filename)
        return ""
