import io

from app.services.extraction import document_parser


def test_extract_plain_text_utf8():
    content = "轴承磨损导致异常振动。".encode("utf-8")
    assert document_parser.extract_text(content, "note.txt") == "轴承磨损导致异常振动。"


def test_extract_plain_text_gbk_fallback():
    content = "冷却器属于液压系统。".encode("gbk")
    result = document_parser.extract_text(content, "note.txt")
    assert "冷却器" in result


def test_extract_docx_paragraphs_and_tables():
    import docx

    document = docx.Document()
    document.add_paragraph("轴承磨损导致异常振动。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "部件"
    table.rows[0].cells[1].text = "轴承"
    buffer = io.BytesIO()
    document.save(buffer)

    result = document_parser.extract_text(buffer.getvalue(), "report.docx")
    assert "轴承磨损导致异常振动" in result
    assert "部件" in result and "轴承" in result


def test_unsupported_extension_returns_empty_not_error():
    assert document_parser.extract_text(b"binary junk", "scan.doc") == ""


def test_corrupt_pdf_returns_empty_not_raises():
    assert document_parser.extract_text(b"not a real pdf", "broken.pdf") == ""
